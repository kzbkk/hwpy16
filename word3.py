''') Создайте абзац с текстом и запишите это в новый word-файл, изменит шрифт и размер
шрифта.'''
import docx
from docx.shared import Pt
doc = docx.Document()
text = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed non risus. Suspendisse lectus tortor, dignissim sit amet, adipiscing nec, ultricies sed, dolor. Cras elementum ultrices diam. Maecenas ligula massa, varius a, semper congue, euismod non, mi. Proin porttitor, orci nec nonummy molestie, enim est eleifend mi, non fermentum diam nisl sit amet erat."
par1 = doc.add_paragraph(text)
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.size = Pt(20)
        run.font.name = 'Times New Roman'
doc.save('task3.docx')