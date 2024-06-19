'''а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран
'''
import docx
doc = docx.Document()
par1 = doc.add_paragraph('Hello ')
par1.add_run('Python').bold = True
print(doc.paragraphs[0].runs[1].text)

doc.save('hellopython.docx')